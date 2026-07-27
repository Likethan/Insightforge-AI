import streamlit as st # type: ignore
import time
import os
import sys
import plotly.express as px # type: ignore
import pandas as pd # type: ignore
import plotly.io as pio # type: ignore
import docx # type: ignore
from io import BytesIO
import base64
import datetime

sys.path.append(os.path.dirname(os.path.abspath(__file__)))

import src.utils.db as db
db.init_db()

# -----------------
# Utility Functions
# -----------------
def get_base64_image(image_path):
    try:
        with open(image_path, "rb") as img_file:
            return base64.b64encode(img_file.read()).decode()
    except Exception as e:
        return None

def create_docx(text):
    doc = docx.Document()
    lines = text.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), 0)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), 1)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), 2)
        else:
            doc.add_paragraph(line)
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def get_mock_report(topic):
    topic_lower = topic.lower()
    if "ev" in topic_lower or "battery" in topic_lower or "charging" in topic_lower:
        return """# Title: Electric Vehicle Charging Grid Constraints 2026

## 1. Executive Summary
The rapid transition towards high-voltage electric vehicle fleets has placed unprecedented demands on regional utility grids. This executive brief explores the bottlenecks in primary substation transmission capacities, local distribution transformers, and the grid enhancements necessary to support 2026 fleet targets.

## 2. Strategic Analysis
The core challenges lie in peak load synchronization. Fleet charging typically surges between 5 PM and 9 PM, aligning with peak residential loads. Solid-state battery implementations promise faster charging times, but require extreme power draws (up to 350kW per vehicle) which local grids cannot support without decentralized buffer storage.

## 3. AI Insights
Deploying edge artificial intelligence at utility nodes can predict load curves with 94.2% accuracy. Active load shaping models can intelligently throttle vehicle charge speeds based on active transformer temperatures and household load forecasts, averting grid collapse and optimizing utility margins.

## 4. Global Trends
European grids have implemented dynamic pricing models, showing a 30% shift in charging schedules to off-peak night hours. Meanwhile, in North America, investment has shifted towards mega-charger depots equipped with integrated solar arrays and battery energy storage systems (BESS).

## 5. Case Studies
- **Vattenfall Smart Grid Initiative (Sweden)**: Leveraged local BESS to buffer peak load draws from rapid bus transport charging depots.
- **ChargePoint Enterprise Pilot (California)**: Tested dynamic demand throttling across 1,200 commercial chargers, cutting peak grid demand by 22%.

## 6. Data Visualization Insights
Key metrics track the correlation between grid capacity saturation and EV penetration rates. Our models project that without active BESS integration, primary substations in urban zones will exceed maximum threshold limits by Q4 2026.

## 7. Conclusion
Decentralized battery storage systems combined with predictive AI scheduling represent the only viable path to support next-generation charging infrastructure without multi-billion dollar grid reconstructions.

## 8. References
1. Global EV Outlook 2026 - International Energy Agency (IEA).
2. Grid Infrastructure and Smart Charging Protocols - IEEE Transactions on Smart Grid, 2025.
"""
    elif "semiconductor" in topic_lower or "fab" in topic_lower or "microchip" in topic_lower:
        return """# Title: Semiconductor Geopolitics: Sub-10nm Expansion

## 1. Executive Summary
Global microchip manufacturing has reached a critical strategic junction as fabrication facilities below the 10nm node scale outside traditional East Asian hubs. This report reviews geopolitical incentives, fab construction timelines, and supply chain dependencies for 2026.

## 2. Strategic Analysis
Building sub-10nm foundries requires specialized Extreme Ultraviolet (EUV) lithography systems, produced exclusively by ASML in the Netherlands. Geopolitical restrictions on the export of these systems create severe supply chain bottlenecks, causing fab construction delays of up to 18 months in newly established regions.

## 3. AI Insights
Predictive supply chain intelligence identifies over 40 single-source critical minerals (e.g., Gallium, Germanium) vulnerable to export blockages. Machine learning models recommend dynamic dual-sourcing options and identify synthetic alternative materials to protect raw supply pipelines.

## 4. Global Trends
The US and European Chips Acts have injected over $80B in direct subsidies, prompting major foundry developments in Arizona, Ohio, and Germany. However, finding and training specialized fabrication engineers remains a critical bottleneck.

## 5. Case Studies
- **TSMC Fab 21 (Arizona)**: Faced construction delay and cultural integration hurdles, shifting its commercial production targets to late 2025/2026.
- **Intel Fab 34 (Ireland)**: Successfully ramped EUV manufacturing under European Chips Act subsidies, serving as a template for Western manufacturing.

## 6. Data Visualization Insights
Data trends analyze sub-10nm global supply share from 2020 through 2026, projecting a gradual decrease in East Asian capacity concentration from 92% to 78% as Western fabs come online.

## 7. Conclusion
While Western subsidies have successfully initiated foundry decentralization, true supply chain independence requires building local chemical processing and lithography tooling networks, which will take another decade.

## 8. References
1. Geopolitical Re-shoring in Global Foundries - McKinsey Semiconductor Report, 2026.
2. The Critical Materials Supply Risk Matrix - US Dept of Energy, 2025.
"""
    elif "agriculture" in topic_lower or "agro" in topic_lower or "farming" in topic_lower or "climate" in topic_lower:
        return """# Title: Sustainable Agro-Tech Trends: Climate Resilience

## 1. Executive Summary
Climate variability in Southern India has driven a rapid transition toward automated climate-resilient farming techniques. This study tracks soil moisture indexing, automated drip-irrigation networks, and drought-resistant crop selections for the 2026 season.

## 2. Strategic Analysis
Erratic monsoon patterns and declining water tables in Tamil Nadu demand a shift from traditional flood irrigation. Smart agro-tech systems leverage IoT soil sensors to coordinate precision drip irrigation, reducing water consumption by 40% while preserving crop yield.

## 3. AI Insights
Satellite crop monitoring models can identify early signs of heat stress 10 days before visible leaf damage. By linking this data to automated micro-fertilizer dispensers, yields can be sustained even during prolonged dry spells.

## 4. Global Trends
- Drip-irrigation systems integrated with solar-powered pump networks are expanding across dryland regions globally, including Sub-Saharan Africa and the Mediterranean, backed by international climate funding.

## 5. Case Studies
- **Coimbatore District IoT Pilot**: Outfitted 50 paddy fields with real-time moisture monitoring, showing a 15% crop yield increase during a water-deficient monsoon.
- **Cauvery Basin Crop Rotation**: Transitioned 2,000 hectares from sugarcane to high-yield drought-resistant millets, reducing local water draw by 50%.

## 6. Data Visualization Insights
Visual graphs map water consumption volumes against crop yield rates, demonstrating the high resource efficiency of automated agro-tech versus manual farming models.

## 7. Conclusion
Widespread adoption of precision IoT agro-tech combined with crop rotation shifts is critical to future-proofing agricultural output against global warming trends.

## 8. References
1. Climate Resilience in Indian Agriculture - Indian Council of Agricultural Research (ICAR).
2. Automated Irrigation Optimization Models - Journal of Agricultural Water Management, 2025.
"""
    else:
        return f"""# Title: Strategic Intelligence Briefing: {topic}

## 1. Executive Summary
This document compiles autonomous strategic intelligence metrics regarding the topic "{topic}". It highlights core structural drivers, technological adoption rates, and operational trends for 2026.

## 2. Strategic Analysis
Initial analysis suggests this topic represents a major area of consolidation. Organizations adopting predictive frameworks are showing faster cycle times and reduced overhead. Key bottlenecks center on technology integration barriers and raw material procurement.

## 3. AI Insights
Predictive neural nets highlight three distinct areas for automation potential, targeting a 12% improvement in throughput over the next 18 months.

## 4. Global Trends
Market shares are decentralizing, moving from centralized manufacturing toward multi-hub distribution models supported by public infrastructure subsidies.

## 5. Case Studies
- **Enterprise Alpha Deployment**: Achieved a 15% cost reduction within 6 months of integrating automated monitoring tools.
- **Continental Logistics Group**: Successfully buffered distribution delays using predictive route optimization algorithms.

## 6. Data Visualization Insights
Data trends project a steady 18% growth rate in adoption across primary markets, with saturation expected in Q3 2027.

## 7. Conclusion
Strategic optimization remains the primary recommendation. Teams should initiate early tooling audits to build technical debt resilience.

## 8. References
1. Technology and Innovation Forecast - McKinsey Global Institutes, 2026.
2. Operations Research Annual Journal - INFORMS, 2025.
"""

# Load LangGraph backend
try:
    from src.graph.workflow import app_graph # type: ignore
except Exception as e:
    st.error(f"Error loading LangGraph backend: {str(e)}")
    app_graph = None

# -----------------
# Page Config & Styles
# -----------------
st.set_page_config(
    page_title="InsightForge AI",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize Session States
if "view" not in st.session_state:
    st.session_state.view = "Workspace"
if "pending_query" not in st.session_state:
    st.session_state.pending_query = ""
if "trigger_research" not in st.session_state:
    st.session_state.trigger_research = False
if "processing" not in st.session_state:
    st.session_state.processing = False
if "sidebar_tab" not in st.session_state:
    st.session_state.sidebar_tab = "Dashboard"

# Initialize Settings States
if "settings_max_iterations" not in st.session_state:
    st.session_state.settings_max_iterations = 2
if "settings_llm_model" not in st.session_state:
    st.session_state.settings_llm_model = "gemini-2.5-pro"
if "settings_web_search" not in st.session_state:
    st.session_state.settings_web_search = True
if "settings_vision_core" not in st.session_state:
    st.session_state.settings_vision_core = True

def load_design_system():
    css_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets", "style.css")
    if os.path.exists(css_path):
        with open(css_path, "r", encoding="utf-8") as f:
            css_data = f.read()
            st.markdown(f"<style>{css_data}</style>", unsafe_allow_html=True)

load_design_system()

# -----------------
# 3-Column Layout Construction (Consolidated Sidebar & Main Dashboard)
# -----------------

# 1. SIDEBAR (Left Column Navigation Rail / Taskbar)
st.sidebar.markdown("""
    <div class="tnra-logo-container">
        <div class="tnra-logo-title">
            <span>⚡</span> InsightForge AI
        </div>
        <div class="tnra-logo-subtitle">Strategic Intelligence Platform</div>
    </div>
""", unsafe_allow_html=True)

# Navigation Rails / Buttons
nav_tabs = [
    ("Dashboard", "🖥️ Dashboard"),
    ("History", "🕒 History"),
    ("Settings", "⚙️ Settings")
]

for tab_id, tab_label in nav_tabs:
    is_active = st.session_state.sidebar_tab == tab_id
    # Add a visual indicator to the label text for additional accessibility
    display_label = f"● {tab_label}" if is_active else f"  {tab_label}"
    key_suffix = "active" if is_active else "inactive"
    
    if st.sidebar.button(display_label, key=f"sidebar_tab_btn_{tab_id}_{key_suffix}", use_container_width=True):
        st.session_state.sidebar_tab = tab_id
        st.rerun()

st.sidebar.markdown("<hr style='border-color: var(--border-color); margin: 15px 0 20px 0;'>", unsafe_allow_html=True)

# Render Content based on Active Tab
if st.session_state.sidebar_tab == "Dashboard":
    # New Session Button
    if st.sidebar.button("✨ New Research Session", key="sidebar_btn_new", use_container_width=True):
        # Clear research states
        st.session_state.current_topic = ""
        st.session_state.pop("report", None)
        st.session_state.pop("final_topic", None)
        st.session_state.pop("visuals", None)
        st.session_state.processing = False
        st.rerun()
        
    st.sidebar.markdown("""
        <div style="font-family: var(--font-family-mono); font-size: 0.72rem; color: var(--text-muted); margin-top: 30px; line-height: 1.6;">
            <div style="margin-bottom: 6px;"><span style="color: var(--accent-primary);">●</span> Engine Status: Ready</div>
            <div style="margin-bottom: 6px;"><span style="color: var(--accent-primary);">●</span> Version: v1.2.0</div>
            <div><span style="color: var(--accent-primary);">●</span> Workspace: Local</div>
        </div>
    """, unsafe_allow_html=True)

elif st.session_state.sidebar_tab == "History":
    st.sidebar.markdown("<div style='font-size:0.75rem; font-weight:600; text-transform:uppercase; color:var(--text-muted); margin-bottom:10px;'>Past Sessions</div>", unsafe_allow_html=True)
    try:
        history_items = db.get_all_runs()
    except Exception as db_err:
        history_items = []
        st.sidebar.error(f"Error loading history: {db_err}")
        
    if not history_items:
        st.sidebar.info("No past research sessions found.")
    else:
        for h_item in history_items:
            col_btn, col_del = st.sidebar.columns([0.82, 0.18])
            
            with col_btn:
                # Truncate title if extremely long for sidebar layout
                display_title = h_item['topic']
                if len(display_title) > 28:
                    display_title = display_title[:25] + "..."
                    
                if st.button(f"📊 {display_title}", key=f"hist_btn_{h_item['id']}", use_container_width=True, help=h_item['topic']):
                    full_run = db.get_run_by_id(h_item['id'])
                    if full_run:
                        st.session_state["report"] = full_run["report"]
                        st.session_state["final_topic"] = full_run["topic"]
                        st.session_state["visuals"] = full_run["visuals"]
                        st.session_state.current_topic = full_run["query"]
                        st.session_state.processing = False
                        # Swap back to Dashboard tab automatically upon loading history
                        st.session_state.sidebar_tab = "Dashboard"
                        st.rerun()
            
            with col_del:
                if st.button("🗑️", key=f"del_btn_{h_item['id']}", help="Delete this session", use_container_width=True):
                    db.delete_run(h_item['id'])
                    st.rerun()
                    
        st.sidebar.markdown("<hr style='border-color: var(--border-color); margin: 15px 0 10px 0;'>", unsafe_allow_html=True)
        if st.sidebar.button("🧹 Clear All History", key="clear_all_history_btn", use_container_width=True):
            db.clear_all_runs()
            st.session_state.pop("report", None)
            st.session_state.pop("final_topic", None)
            st.session_state.pop("visuals", None)
            st.session_state.current_topic = ""
            st.rerun()

elif st.session_state.sidebar_tab == "Settings":
    st.sidebar.markdown("<div style='font-size:0.75rem; font-weight:600; text-transform:uppercase; color:var(--text-muted); margin-bottom:10px;'>Parameters</div>", unsafe_allow_html=True)
    st.session_state.settings_llm_model = st.sidebar.selectbox(
        "AI Model (Intelligence)",
        ["gemini-2.5-pro", "gemini-2.5-flash", "gemini-2.0-flash"],
        index=0,
        key="model_select_box"
    )
    st.session_state.settings_max_iterations = st.sidebar.slider(
        "Max Iterations (Critique Loops)",
        min_value=1, max_value=5, value=st.session_state.settings_max_iterations,
        key="iterations_slider"
    )
    st.session_state.settings_language = st.sidebar.selectbox(
        "Language Focus",
        ["English", "Tamil", "Hindi"],
        index=0,
        key="language_select_box"
    )
    
    st.sidebar.markdown("<div style='height: 15px;'></div>", unsafe_allow_html=True)
    
    # Export Settings Button
    if st.sidebar.button("🖨️ Export All Settings", key="sidebar_btn_export", use_container_width=True):
        st.sidebar.success("Settings configuration exported!")

# Check state to decide layout
has_report = "report" in st.session_state
is_processing = st.session_state.processing

col_center = None
col_right = None
pipeline_placeholder = None
agent_status_placeholder = None
report_preview_placeholder = None

start_research = False
query = ""

if not has_report and not is_processing:
    # -----------------
    # LANDING STATE (Idle, Centered UI)
    # -----------------
    # Load base64 logo
    logo_base64 = get_base64_image(os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets", "logo.png"))
    
    # HTML structure for centered logo, title, and subtitle
    logo_html = ""
    if logo_base64:
        logo_html = f'''<div class="landing-logo-container">
<div class="logo-box">
<img class="logo-img" src="data:image/png;base64,{logo_base64}" />
</div>
</div>'''
    else:
        # Fallback if logo not found
        logo_html = '''<div class="landing-logo-container">
<div class="logo-box">
<span style="font-size: 2.5rem;">⚡</span>
</div>
</div>'''

    st.markdown(f'''<div class="landing-container">
{logo_html}
<h1 class="landing-title">InsightForge AI</h1>
<div class="landing-subtitle">Autonomous Strategic Intelligence Platform</div>
</div>''', unsafe_allow_html=True)
    
    # Query Search Box Container (Centered)
    _, col_input, _ = st.columns([1, 6, 1])
    with col_input:
        st.markdown('<div class="landing-search-container">', unsafe_allow_html=True)
        query_landing = st.text_input(
            "Research Topic",
            placeholder="Ask InsightForge AI to research anything...",
            label_visibility="collapsed",
            key="landing_topic_input"
        )
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Button Row
        start_research_landing = st.button("INITIALIZE ENGINE 🚀", key="btn_initialize_engine", use_container_width=True)
                
        # Centered Agent Chips at the bottom
        st.markdown('''
            <div class="agent-chips-container">
                <span class="agent-chip"><span class="chip-dot"></span>Supervisor</span>
                <span class="agent-chip"><span class="chip-dot"></span>Researcher</span>
                <span class="agent-chip"><span class="chip-dot"></span>Analyzer</span>
                <span class="agent-chip"><span class="chip-dot"></span>Writer</span>
            </div>
        ''', unsafe_allow_html=True)
        
    # Map triggers
    query = query_landing
    start_research = start_research_landing

    # To catch Enter key triggers:
    if query_landing and query_landing != st.session_state.get("current_topic", ""):
        st.session_state.processing = True
        st.session_state.current_topic = query_landing
        st.session_state.pop("report", None)
        st.session_state.pop("final_topic", None)
        st.session_state.pop("visuals", None)
        st.rerun()

else:
    # -----------------
    # ACTIVE WORKSPACE STATE (Split 2-Column Dashboard & Report Preview)
    # -----------------
    col_center, col_right = st.columns([1.1, 0.9])
    
    with col_center:
        st.markdown('<div class="section-title" style="margin-top:0;">InsightForge Research Dashboard</div>', unsafe_allow_html=True)
        
        # Topic query input
        default_query = st.session_state.get("current_topic", "")
        st.markdown('<div class="landing-search-container">', unsafe_allow_html=True)
        query_active = st.text_input(
            "Research Topic",
            value=default_query,
            placeholder="Ask InsightForge AI to research anything...",
            label_visibility="collapsed",
            key="research_topic_input"
        )
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Button Row
        start_research_active = st.button("INITIALIZE ENGINE 🚀", key="btn_initialize_engine_active", use_container_width=True)
                
        st.markdown("<div style='height: 15px;'></div>", unsafe_allow_html=True)
        
        # Dynamic Placeholders for Phase Status and Agent status
        pipeline_placeholder = st.empty()
        st.markdown('<h3 style="font-size:1.15rem; font-weight:700; color:#ffffff; margin-bottom: 10px;">🤖 Multi-Agent Team Status</h3>', unsafe_allow_html=True)
        agent_status_placeholder = st.empty()
        st.markdown("<p style='font-size:0.8rem; color:#64748b; cursor:pointer; margin-top: 5px;'>View Agent Logs</p>", unsafe_allow_html=True)
        
    with col_right:
        st.markdown('<div class="section-title" style="margin-top:0;">📄 Live Report Preview</div>', unsafe_allow_html=True)
        st.markdown('<div style="font-size:0.82rem; color:#64748b; text-transform:uppercase; font-weight:700; margin-top:-10px; margin-bottom:20px;">Generated by InsightForge AI</div>', unsafe_allow_html=True)
        
        # Dynamic Placeholder for live report preview
        report_preview_placeholder = st.empty()
        
    # Map triggers
    query = query_active
    start_research = start_research_active
    
    # To catch Enter key triggers:
    if query_active and query_active != st.session_state.get("current_topic", ""):
        st.session_state.processing = True
        st.session_state.current_topic = query_active
        st.session_state.pop("report", None)
        st.session_state.pop("final_topic", None)
        st.session_state.pop("visuals", None)
        st.rerun()

# -----------------
# Execution Routing and Rendering
# -----------------

# Helper function to generate agent status HTML cards
def get_agent_cards_html(active_agent=None, completed_agents=[]):
    agents = [
        {"name": "Supervisor", "icon": "👑"},
        {"name": "Researcher", "icon": "🔎"},
        {"name": "Analyzer", "icon": "⚙️"},
        {"name": "Critic", "icon": "🛡️"},
        {"name": "Visualizer", "icon": "📊"},
        {"name": "Writer", "icon": "✍️"}
    ]
    
    html = '<div style="display:flex; flex-direction:column; gap:4px;">'
    for agent in agents:
        name = agent["name"]
        icon = agent["icon"]
        if active_agent == name:
            status = "Working"
            status_class = "working"
        elif name in completed_agents:
            status = "Completed"
            status_class = "completed"
        else:
            status = "Idle"
            status_class = "idle"
            
        html += f'<div class="agent-card"><div class="agent-card-info"><span class="agent-card-icon">{icon}</span><span class="agent-card-name">{name} Agent</span></div><span class="agent-card-status {status_class}">{status}</span></div>'
    html += '</div>'
    return html

# Trigger Research Execution
if start_research and query:
    st.session_state.processing = True
    st.session_state.current_topic = query
    st.session_state.pop("report", None)
    st.session_state.pop("final_topic", None)
    st.session_state.pop("visuals", None)
    st.rerun()

# Processing State Loop
if st.session_state.processing:
    # Ensure placeholders are initialized (for type checking & safe rendering)
    if (pipeline_placeholder is not None 
        and agent_status_placeholder is not None 
        and report_preview_placeholder is not None
        and col_center is not None):
        
        # Setup initial pipeline status
        pipeline_placeholder.markdown("""
            <div style="margin-bottom: 20px;">
                <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom: 6px;">
                    <span style="font-size:0.85rem; color:#94a3b8; font-weight:600;">Current Phase: Initializing Supervisor</span>
                    <span style="font-size:0.85rem; color:#3b82f6; font-weight:700;">10%</span>
                </div>
                <div style="width: 100%; background: #27272a; height: 6px; border-radius: 4px; overflow: hidden; margin-bottom: 8px;">
                    <div style="width: 10%; background: #3b82f6; height: 100%; border-radius: 4px;"></div>
                </div>
            </div>
        """, unsafe_allow_html=True)
        
        agent_status_placeholder.markdown(get_agent_cards_html("Supervisor", []), unsafe_allow_html=True)
        report_preview_placeholder.markdown("<div style='color: #64748b; font-style: italic; padding: 20px 0;'>Awaiting agent pipeline to start generation...</div>", unsafe_allow_html=True)
        
        # Cancel/Pause buttons
        col_pause, col_cancel = col_center.columns(2)
        with col_pause:
            st.button("⏸️ Pause", key="btn_pause_run", use_container_width=True)
        with col_cancel:
            if st.button("🔴 Cancel", key="btn_cancel_run", use_container_width=True):
                st.session_state.processing = False
                st.rerun()

        initial_state = {
            "topic": st.session_state.current_topic,
            "max_iterations": st.session_state.settings_max_iterations,
            "iterations": 0,
            "llm_model": st.session_state.settings_llm_model
        }
        
        if not app_graph:
            st.error("Intelligence Backend is offline.")
            st.session_state.processing = False
            st.rerun()
        else:
            try:
                final_state = {}
                completed = []
                
                def update_ui_stream(percent, phase_label, active_agent, completed_agents, partial_report=None):
                    if pipeline_placeholder is not None and agent_status_placeholder is not None and report_preview_placeholder is not None:
                        pipeline_placeholder.markdown(f"""
                            <div style="margin-bottom: 20px;">
                                <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom: 6px;">
                                    <span style="font-size:0.85rem; color:#94a3b8; font-weight:600;">Current Phase: {phase_label}</span>
                                    <span style="font-size:0.85rem; color:#3b82f6; font-weight:700;">{percent}%</span>
                                </div>
                                <div style="width: 100%; background: #27272a; height: 6px; border-radius: 4px; overflow: hidden; margin-bottom: 8px;">
                                    <div style="width: {percent}%; background: #3b82f6; height: 100%; border-radius: 4px; transition: width 0.4s ease;"></div>
                                </div>
                            </div>
                        """, unsafe_allow_html=True)
                        agent_status_placeholder.markdown(get_agent_cards_html(active_agent, completed_agents), unsafe_allow_html=True)
                        if partial_report:
                            report_preview_placeholder.markdown(partial_report)
                        else:
                            report_preview_placeholder.markdown(f"<div style='color: #64748b; font-style: italic; padding: 20px 0;'>Generating content via {active_agent} Agent...</div>", unsafe_allow_html=True)
                
                # Execute LangGraph Stream
                for output in app_graph.stream(initial_state):
                    for node_name, state in output.items():
                        if node_name == "supervisor":
                            completed.append("Supervisor")
                            update_ui_stream(15, "Supervisor Planning", "Researcher", completed)
                        elif node_name == "researcher":
                            completed.append("Researcher")
                            update_ui_stream(35, "Scanning Web Sources", "Analyzer", completed)
                        elif node_name == "analyzer":
                            completed.append("Analyzer")
                            update_ui_stream(60, "Analyzing Context", "Critic", completed)
                        elif node_name == "critic":
                            completed.append("Critic")
                            update_ui_stream(80, "Validating Strategic Authenticity", "Visualizer", completed)
                        elif node_name == "visualizer":
                            completed.append("Visualizer")
                            update_ui_stream(90, "Plotting Data Charts", "Writer", completed)
                        elif node_name == "writer":
                            completed.append("Writer")
                            report_text = state.get("report", "Report compiled successfully.")
                            update_ui_stream(100, "Writing Report", None, completed, partial_report=report_text)
                        final_state = state
                
                st.session_state["report"] = final_state.get("report", "No report generated.")
                st.session_state["visuals"] = final_state.get("visuals", {})
                st.session_state["final_topic"] = st.session_state.current_topic
                
                # Save run to SQLite history database
                try:
                    db.save_run(
                        topic=st.session_state.current_topic,
                        query=st.session_state.current_topic,
                        report=st.session_state["report"],
                        visuals=st.session_state["visuals"]
                    )
                except Exception as save_err:
                    st.warning(f"Error saving to history: {save_err}")
                
                st.session_state.processing = False
                time.sleep(0.5)
                st.rerun()
                
            except Exception as e:
                st.error(f"Critical System Error: {str(e)}")
                st.session_state.processing = False
                st.rerun()

# Completed / Idle State Layout
else:
    # If report exists, display complete details
    if "report" in st.session_state:
        # Ensure placeholders are initialized (satisfy static analysis type checker)
        if (pipeline_placeholder is not None 
            and agent_status_placeholder is not None 
            and report_preview_placeholder is not None):
            
            # Success message
            pipeline_placeholder.markdown("""
                <div style="margin-bottom: 20px;">
                    <div style="font-size:0.95rem; color:#10b981; font-weight:700; display:flex; align-items:center; gap:8px;">
                        <span>✓</span> Research Completed Successfully! 🎉
                    </div>
                </div>
            """, unsafe_allow_html=True)
            
            # Display agent status cards as all Completed
            agent_status_placeholder.markdown(get_agent_cards_html(None, ["Supervisor", "Researcher", "Analyzer", "Critic", "Visualizer", "Writer"]), unsafe_allow_html=True)
            
            # Display report content & downloads in right column
            report_text = st.session_state["report"]
            sections = report_text.split("##")
            
            with report_preview_placeholder.container():
                # Action download bar
                c1, c2, c3 = st.columns([1, 1, 1])
                with c1:
                    docx_data = create_docx(report_text)
                    st.download_button(label="Download DOCX", data=docx_data, file_name=f"Report_{int(time.time())}.docx", use_container_width=True)
                with c2:
                    st.button("Export PDF", key="btn_pdf_dl", use_container_width=True)
                with c3:
                    st.button("Copy Text", key="btn_copy_txt", use_container_width=True)
                    
                st.markdown("<div style='height: 15px;'></div>", unsafe_allow_html=True)
                st.markdown(f"<div style='font-family: var(--font-family); font-weight: 700; font-size: 1.6rem; color: #ffffff; margin-bottom: 1.5rem;'>Research Report: {st.session_state.get('final_topic', 'Research Report')}</div>", unsafe_allow_html=True)
                
                # Render report sections
                if len(sections) > 1:
                    for section in sections[1:]:
                        lines = section.strip().split("\n")
                        header = lines[0].strip()
                        clean_header = header.lstrip("1234567890. ")
                        body = "\n".join(lines[1:]).strip()
                        
                        with st.container():
                            st.markdown("<div class='report-section-card-marker'></div>", unsafe_allow_html=True)
                            st.markdown(f"<div class='report-section-header'>{clean_header}</div>", unsafe_allow_html=True)
                            st.markdown(f"<div class='report-document-body'>\n\n{body}\n\n</div>", unsafe_allow_html=True)
                            
                            # Inline Plotly Chart rendering for Projections / Visualizations section
                            if "PROJECTIONS" in clean_header.upper() or "VISUALIZATION" in clean_header.upper() or "CHART" in clean_header.upper():
                                visuals = st.session_state.get("visuals", {})
                                if visuals and isinstance(visuals, dict) and "data" in visuals:
                                    st.markdown("<div style='height: 10px;'></div>", unsafe_allow_html=True)
                                    try:
                                        minimalist_layout = dict(
                                            plot_bgcolor='rgba(0,0,0,0)',
                                            paper_bgcolor='rgba(0,0,0,0)',
                                            font=dict(color='#a1a1aa', family='Inter'),
                                            xaxis=dict(showgrid=False, zeroline=False),
                                            yaxis=dict(gridcolor='rgba(255,255,255,0.05)', zeroline=False)
                                        )
                                        if visuals.get("is_json_string") and isinstance(visuals["data"], str):
                                            fig = pio.from_json(visuals["data"])
                                            fig.update_layout(**minimalist_layout)
                                            st.plotly_chart(fig, use_container_width=True)
                                        elif "data" in visuals:
                                            df = pd.DataFrame(visuals["data"])
                                            if not df.empty and "x" in df.columns and "y" in df.columns:
                                                fig = px.bar(df, x="x", y="y", title=visuals.get("title", ""), template='plotly_dark')
                                                fig.update_layout(**minimalist_layout)
                                                st.plotly_chart(fig, use_container_width=True)
                                            else:
                                                st.table(df)
                                    except Exception as chart_err:
                                        st.warning(f"Could not render chart: {chart_err}")
                else:
                    with st.container():
                        st.markdown("<div class='report-section-card-marker'></div>", unsafe_allow_html=True)
                        st.markdown(f"<div class='report-document-body'>\n\n{report_text}\n\n</div>", unsafe_allow_html=True)
    else:
        # We are in the landing state, nothing to do as the centered UI is already drawn
        pass